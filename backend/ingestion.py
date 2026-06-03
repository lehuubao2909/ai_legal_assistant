import os
import re
import json
import shutil
from datetime import datetime
from docx import Document
from typing import List, Dict, Any, Optional
from html import unescape
import chromadb

class LegalDocumentIngestion:
    def __init__(self, db_path: str = None, registry_path: str = None):
        current_dir = os.path.dirname(os.path.abspath(__file__))
        
        if registry_path is None:
            registry_path = os.path.join(current_dir, "..", "data", "document_registry.json")
            
        self.registry_path = os.path.abspath(registry_path)
        
        self.uploaded_docs_dir = os.path.join(current_dir, "static", "uploaded_documents")
        os.makedirs(self.uploaded_docs_dir, exist_ok=True)
        
        # Connect to ChromaDB
        chroma_path = os.path.join(current_dir, "..", "data", "chroma_db")
        self.chroma_client = chromadb.PersistentClient(path=chroma_path)
        self.collection = self.chroma_client.get_or_create_collection(
            name="vietnamese_laws",
            metadata={"hnsw:space": "cosine"}
        )

    def _generate_embedding(self, text: str) -> Optional[List[float]]:
        """Calls Gemini API (google-genai SDK) to generate real semantic vector embeddings."""
        from google import genai as google_genai
        from google.genai import types as genai_types
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            print("Warning: GEMINI_API_KEY not set. Embedding will be default zero vector.")
            return None
        try:
            client = google_genai.Client(api_key=api_key)
            result = client.models.embed_content(
                model="gemini-embedding-001",
                contents=[text],
                config=genai_types.EmbedContentConfig(task_type="RETRIEVAL_DOCUMENT"),
            )
            if result.embeddings:
                return list(result.embeddings[0].values)
            return None
        except Exception as e:
            print(f"Error generating embedding during ingestion: {e}")
            return None

    def _load_existing_db(self) -> List[Dict[str, Any]]:
        """Dummy method for backward compatibility."""
        return []

    def _save_db(self, data: List[Dict[str, Any]]):
        """Dummy method for backward compatibility."""
        pass

    def _load_registry(self) -> List[Dict[str, Any]]:
        if not os.path.exists(self.registry_path):
            return []
        try:
            with open(self.registry_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return []

    def _save_registry(self, data: List[Dict[str, Any]]):
        with open(self.registry_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    def clean_html_tags(self, html_text: str) -> str:
        """Helper to strip HTML tags and decode unescaped characters."""
        text = re.sub(r'<[^>]*>', '', html_text)
        return unescape(text).strip()

    def parse_docx(self, docx_path: str, law_name_override: Optional[str] = None, source_url: Optional[str] = None) -> Dict[str, Any]:
        """Parses a raw Vietnamese legal .docx file into structured parent-child JSON rules."""
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Không tìm thấy file: {docx_path}")
            
        doc = Document(docx_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # 1. Detect Law Name
        law_name = law_name_override
        if not law_name:
            for p in paragraphs[:10]:
                if any(k in p.upper() for k in ["LUẬT", "BỘ LUẬT", "NGHỊ ĐỊNH", "THÔNG TƯ", "QUYẾT ĐỊNH"]):
                    law_name = p
                    break
            if not law_name:
                law_name = os.path.basename(docx_path).replace(".docx", "").replace("_", " ")
        
        # Default Metadata
        effective_date = datetime.now().strftime("%Y-%m-%d")
        for p in paragraphs[:15] + paragraphs[-10:]:
            match_date = re.search(r"ngày\s+(\d+)\s+tháng\s+(\d+)\s+năm\s+(\d{4})", p.lower())
            if match_date:
                day, month, year = match_date.groups()
                effective_date = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
                break

        # Abbreviate law name for unique IDs
        law_abbr = re.sub(r"[^a-zA-Z0-9\s]", "", law_name)
        law_abbr = "".join([w[0].upper() for w in law_abbr.split() if w])
        match_year = re.search(r"\b(20\d{2}|19\d{2})\b", law_name)
        if match_year:
            law_abbr += f"_{match_year.group(1)}"
        else:
            law_abbr += f"_{datetime.now().year}"

        parsed_items = []
        current_parent = None
        
        article_regex = re.compile(r"^Điều\s+(\d+)\.\s*(.*)", re.IGNORECASE)
        clause_regex = re.compile(r"^(\d+)\.\s*(.*)")
        point_regex = re.compile(r"^([a-đg-y])\)\s*(.*)", re.IGNORECASE)

        for p in paragraphs:
            art_match = article_regex.match(p)
            if art_match:
                art_num, art_title = art_match.groups()
                art_id = f"{law_abbr}_D{art_num}"
                
                current_parent = {
                    "id": art_id,
                    "type": "parent",
                    "law_name": law_name,
                    "title": f"Điều {art_num}. {art_title}",
                    "article": f"Điều {art_num}",
                    "clause": "",
                    "point": "",
                    "text": p,
                    "effective_date": effective_date,
                    "expiration_date": "",
                    "status": "active",
                    "source_url": source_url or ""
                }
                parsed_items.append(current_parent)
                continue

            clause_match = clause_regex.match(p)
            if clause_match and current_parent:
                cls_num, cls_text = clause_match.groups()
                cls_id = f"{current_parent['id']}_K{cls_num}"
                
                current_parent["text"] += f"\n{p}"
                
                child_clause = {
                    "id": cls_id,
                    "type": "child",
                    "parent_id": current_parent["id"],
                    "law_name": law_name,
                    "title": f"{current_parent['article']} Khoản {cls_num} - {law_name}",
                    "article": current_parent["article"],
                    "clause": f"Khoản {cls_num}",
                    "point": "",
                    "text": p,
                    "effective_date": effective_date,
                    "expiration_date": "",
                    "status": "active",
                    "source_url": source_url or ""
                }
                parsed_items.append(child_clause)
                continue

            point_match = point_regex.match(p)
            if point_match and current_parent:
                pt_letter, pt_text = point_match.groups()
                pt_id = f"{current_parent['id']}_P{pt_letter.lower()}"
                
                current_parent["text"] += f"\n{p}"
                
                child_point = {
                    "id": pt_id,
                    "type": "child",
                    "parent_id": current_parent["id"],
                    "law_name": law_name,
                    "title": f"{current_parent['article']} Điểm {pt_letter}) - {law_name}",
                    "article": current_parent["article"],
                    "clause": "",
                    "point": f"Điểm {pt_letter})",
                    "text": p,
                    "effective_date": effective_date,
                    "expiration_date": "",
                    "status": "active",
                    "source_url": source_url or ""
                }
                parsed_items.append(child_point)
                continue

            if current_parent and p not in current_parent["text"]:
                current_parent["text"] += f"\n{p}"

        return {
            "law_name": law_name,
            "law_abbr": law_abbr,
            "items": parsed_items
        }

    def ingest_document(self, docx_path: str, law_name_override: Optional[str] = None, source_url: Optional[str] = None) -> int:
        """Parses a docx document, merges items into ChromaDB, saves raw docx physically, and updates the registry."""
        result = self.parse_docx(docx_path, law_name_override, source_url)
        parsed_items = result["items"]
        
        if not parsed_items:
            return 0
            
        law_name = result["law_name"]
        doc_id = result["law_abbr"]
        
        # Archive the raw file permanently
        filename = f"{doc_id}.docx"
        permanent_path = os.path.join(self.uploaded_docs_dir, filename)
        
        if os.path.abspath(docx_path) != os.path.abspath(permanent_path):
            shutil.copy2(docx_path, permanent_path)
            
        file_size_kb = round(os.path.getsize(permanent_path) / 1024, 1)
        total_articles = sum(1 for item in parsed_items if item["type"] == "parent")

        # Update Registry
        registry = self._load_registry()
        registry_map = {doc["doc_id"]: doc for doc in registry}
        
        registry_map[doc_id] = {
            "doc_id": doc_id,
            "filename": filename,
            "law_name": law_name,
            "uploaded_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "file_size_kb": file_size_kb,
            "total_articles": total_articles,
            "source_url": source_url or ""
        }
        self._save_registry(list(registry_map.values()))
            
        # Bulk-generate embeddings and upload to ChromaDB
        ids = []
        embeddings = []
        documents = []
        metadatas = []
        
        print(f"Generating embeddings for {len(parsed_items)} items of '{law_name}'...")
        for item in parsed_items:
            item_id = item["id"]
            content_to_embed = f"{item['title']}\n{item['text']}"
            
            embedding = self._generate_embedding(content_to_embed)
            if not embedding:
                embedding = [0.0] * 768  # Fallback zero-embedding
                
            ids.append(item_id)
            embeddings.append(embedding)
            documents.append(content_to_embed)
            
            # Simple types only in metadata
            metadata = {
                "id": item_id,
                "law_name": item.get("law_name", ""),
                "title": item.get("title", ""),
                "text": item.get("text", ""),
                "type": item.get("type", "parent"),
                "parent_id": item.get("parent_id") or "",
                "effective_date": item.get("effective_date", ""),
                "expiration_date": item.get("expiration_date") or "",
                "status": item.get("status", "active"),
                "source_url": item.get("source_url") or "",
                "article": item.get("article") or "",
                "clause": item.get("clause") or "",
                "point": item.get("point") or ""
            }
            metadatas.append(metadata)
            
        if ids:
            self.collection.upsert(
                ids=ids,
                embeddings=embeddings,
                documents=documents,
                metadatas=metadatas
            )
            
        print(f"Ingested & Archived '{law_name}' (ID: {doc_id}): {total_articles} Articles.")
        return len(parsed_items)

    def ingest_vbpl_payload(self, doc_id: str, payload: Dict[str, Any]) -> int:
        """Parses dynamic HTML response from VBPL, compiles a physical archived docx, and updates registries."""
        data = payload.get("data", {})
        title = data.get("title", f"Văn bản VBPL {doc_id}")
        
        # 1. Parse effective date
        eff_from_raw = data.get("effFrom", "")
        effective_date = datetime.now().strftime("%Y-%m-%d")
        if eff_from_raw:
            try:
                # E.g. "2020-07-01T00:00:00" -> "2020-07-01"
                effective_date = eff_from_raw.split("T")[0]
            except Exception:
                pass
                
        source_url = f"https://vbpl.vn/tw/Pages/vbpq-toanvan.aspx?ItemID={doc_id}"
        
        # 2. Extract and Parse HTML content
        doc_content = data.get("documentContent", {})
        html_content = doc_content.get("content", "")
        
        if not html_content:
            raise ValueError("Không tìm thấy nội dung văn bản (documentContent.content) từ VBPL Gateway.")
            
        # Unique abbreviation ID
        law_abbr = f"VBPL_{doc_id}"
        
        # Find paragraphs inside HTML using regex
        p_tags = re.findall(r'<p([^>]*)>(.*?)</p>', html_content, re.DOTALL | re.IGNORECASE)
        
        parsed_items = []
        current_parent = None
        
        # List of plain text paragraphs to compile into a .docx later
        raw_paragraphs = []

        for attrs, inner_html in p_tags:
            text = self.clean_html_tags(inner_html)
            if not text:
                continue
                
            raw_paragraphs.append(text)
            
            # Detect class from attributes
            class_match = re.search(r'class=["\']([^"\']+)["\']', attrs, re.IGNORECASE)
            p_class = class_match.group(1).lower() if class_match else ""
            
            is_article = "prov-article" in p_class or text.lower().startswith("điều ")
            is_clause = "prov-clause" in p_class or re.match(r"^\d+\.", text)
            is_point = "prov-item" in p_class or re.match(r"^[a-đg-y]\)", text.lower())
            
            if is_article:
                art_num_match = re.search(r"^Điều\s+(\d+)", text, re.IGNORECASE)
                art_num = art_num_match.group(1) if art_num_match else len(parsed_items) + 1
                art_id = f"{law_abbr}_D{art_num}"
                
                current_parent = {
                    "id": art_id,
                    "type": "parent",
                    "law_name": title,
                    "title": text,
                    "article": f"Điều {art_num}",
                    "clause": "",
                    "point": "",
                    "text": text,
                    "effective_date": effective_date,
                    "expiration_date": "",
                    "status": "active",
                    "source_url": source_url
                }
                parsed_items.append(current_parent)
                
            elif is_clause and current_parent:
                cls_num_match = re.match(r"^(\d+)", text)
                cls_num = cls_num_match.group(1) if cls_num_match else "1"
                cls_id = f"{current_parent['id']}_K{cls_num}"
                
                current_parent["text"] += f"\n{text}"
                
                child_clause = {
                    "id": cls_id,
                    "type": "child",
                    "parent_id": current_parent["id"],
                    "law_name": title,
                    "title": f"{current_parent['article']} Khoản {cls_num} - {title}",
                    "article": current_parent["article"],
                    "clause": f"Khoản {cls_num}",
                    "point": "",
                    "text": text,
                    "effective_date": effective_date,
                    "expiration_date": "",
                    "status": "active",
                    "source_url": source_url
                }
                parsed_items.append(child_clause)
                
            elif is_point and current_parent:
                pt_letter_match = re.match(r"^([a-đg-y])", text.lower())
                pt_letter = pt_letter_match.group(1) if pt_letter_match else "a"
                pt_id = f"{current_parent['id']}_P{pt_letter}"
                
                current_parent["text"] += f"\n{text}"
                
                child_point = {
                    "id": pt_id,
                    "type": "child",
                    "parent_id": current_parent["id"],
                    "law_name": title,
                    "title": f"{current_parent['article']} Điểm {pt_letter}) - {title}",
                    "article": current_parent["article"],
                    "clause": "",
                    "point": f"Điểm {pt_letter})",
                    "text": text,
                    "effective_date": effective_date,
                    "expiration_date": "",
                    "status": "active",
                    "source_url": source_url
                }
                parsed_items.append(child_point)
                
            elif current_parent:
                current_parent["text"] += f"\n{text}"

        if not parsed_items:
            raise ValueError("Không thể bóc tách bất kỳ Điều/Khoản/Điểm nào từ mã HTML của VBPL.")

        # 3. Create a physical archived docx file dynamically so user can view/download easily later
        filename = f"{law_abbr}.docx"
        permanent_path = os.path.join(self.uploaded_docs_dir, filename)
        
        doc = Document()
        # Add metadata headers
        doc.add_paragraph("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM").runs[0].bold = True
        doc.add_paragraph("Độc lập - Tự do - Hạnh phúc").runs[0].bold = True
        doc.add_paragraph("-----------------------------------")
        doc.add_paragraph(title).runs[0].bold = True
        doc.add_paragraph(f"Nguồn: Cổng thông tin vbpl.vn • ID: {doc_id}").runs[0].italic = True
        doc.add_paragraph("\n")
        
        # Append all parsed paragraphs
        for p_text in raw_paragraphs:
            doc.add_paragraph(p_text)
            
        doc.save(permanent_path)
        file_size_kb = round(os.path.getsize(permanent_path) / 1024, 1)
        total_articles = sum(1 for item in parsed_items if item["type"] == "parent")

        # 4. Save to Registry
        registry = self._load_registry()
        registry_map = {doc["doc_id"]: doc for doc in registry}
        
        registry_map[law_abbr] = {
            "doc_id": law_abbr,
            "filename": filename,
            "law_name": title,
            "uploaded_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "file_size_kb": file_size_kb,
            "total_articles": total_articles,
            "source_url": source_url
        }
        self._save_registry(list(registry_map.values()))

        # 5. Bulk-generate embeddings and upload to ChromaDB
        ids = []
        embeddings = []
        documents = []
        metadatas = []
        
        print(f"Generating embeddings for {len(parsed_items)} items from VBPL '{title}'...")
        for item in parsed_items:
            item_id = item["id"]
            content_to_embed = f"{item['title']}\n{item['text']}"
            
            embedding = self._generate_embedding(content_to_embed)
            if not embedding:
                embedding = [0.0] * 768
                
            ids.append(item_id)
            embeddings.append(embedding)
            documents.append(content_to_embed)
            
            # Simple types only in metadata
            metadata = {
                "id": item_id,
                "law_name": item.get("law_name", ""),
                "title": item.get("title", ""),
                "text": item.get("text", ""),
                "type": item.get("type", "parent"),
                "parent_id": item.get("parent_id") or "",
                "effective_date": item.get("effective_date", ""),
                "expiration_date": item.get("expiration_date") or "",
                "status": item.get("status", "active"),
                "source_url": item.get("source_url") or "",
                "article": item.get("article") or "",
                "clause": item.get("clause") or "",
                "point": item.get("point") or ""
            }
            metadatas.append(metadata)
            
        if ids:
            self.collection.upsert(
                ids=ids,
                embeddings=embeddings,
                documents=documents,
                metadatas=metadatas
            )

        print(f"Ingested and compiled VBPL document '{title}' (ID: {law_abbr}) with {total_articles} Articles.")
        return len(parsed_items)

    def delete_document(self, doc_id: str) -> bool:
        """Deletes a raw document from storage, clears registry metadata, and deletes its rules from ChromaDB."""
        registry = self._load_registry()
        doc_entry = next((doc for doc in registry if doc["doc_id"] == doc_id), None)
        
        if not doc_entry:
            print(f"Document with ID {doc_id} not found in registry.")
            return False
            
        # 1. Delete physical .docx file
        filename = doc_entry["filename"]
        file_path = os.path.join(self.uploaded_docs_dir, filename)
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception as e:
                print(f"Error deleting physical file {file_path}: {e}")
                
        # 2. Update Registry
        updated_registry = [doc for doc in registry if doc["doc_id"] != doc_id]
        self._save_registry(updated_registry)
        
        # 3. Clean up ChromaDB rules using exact metadata matching
        law_name = doc_entry["law_name"]
        try:
            self.collection.delete(where={"law_name": law_name})
            print(f"Deleted all clauses matching law_name '{law_name}' from ChromaDB.")
        except Exception as e:
            print(f"Error deleting from ChromaDB by law_name: {e}")
            
        print(f"Deleted Document '{law_name}' (ID: {doc_id}) and purged all RAG indices successfully.")
        return True

if __name__ == "__main__":
    ingestor = LegalDocumentIngestion()
    print("Ingestion engine loaded with ChromaDB collection count:", ingestor.collection.count())
