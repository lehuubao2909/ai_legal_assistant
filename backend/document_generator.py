import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any

class LegalDocumentGenerator:
    def __init__(self, output_dir: str = None):
        if output_dir is None:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            output_dir = os.path.join(current_dir, "static", "generated_docs")
        
        self.output_dir = os.path.abspath(output_dir)
        os.makedirs(self.output_dir, exist_ok=True)

    def generate_dismissal_decision(self, data: Dict[str, Any]) -> str:
        """Generates a professional dismissal decision document (.docx)."""
        doc = Document()
        
        # Styles setup
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(13)

        # Header - National Motto
        p_header1 = doc.add_paragraph()
        p_header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_header1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
        run.bold = True
        run.font.size = Pt(13)
        run2 = p_header1.add_run("Độc lập - Tự do - Hạnh phúc\n")
        run2.bold = True
        run2.font.size = Pt(14)
        p_header1.add_run("---------------o0o---------------")

        # Company Title & Date
        company_name = data.get("company_name", "[TÊN DOANH NGHIỆP CỦA BẠN]").upper()
        p_company = doc.add_paragraph()
        run_comp = p_company.add_run(f"ĐƠN VỊ: {company_name}\n")
        run_comp.bold = True
        
        date_str = data.get("date_str", "ngày ... tháng ... năm ...")
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_date.add_run(f"Hà Nội, {date_str}")

        # Document Title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(18)
        p_title.paragraph_format.space_after = Pt(18)
        run_title = p_title.add_run("QUYẾT ĐỊNH\n")
        run_title.bold = True
        run_title.font.size = Pt(16)
        run_sub = p_title.add_run("V/v: Xử lý kỷ luật sa thải người lao động")
        run_sub.bold = True
        run_sub.italic = True

        # Preamble (Căn cứ pháp lý)
        p_base = doc.add_paragraph()
        p_base.add_run("- Căn cứ Bộ luật Lao động nước Cộng hòa xã hội chủ nghĩa Việt Nam năm 2019;\n").italic = True
        p_base.add_run("- Căn cứ Nghị định số 145/2020/NĐ-CP hướng dẫn thi hành Bộ luật Lao động;\n").italic = True
        p_base.add_run("- Căn cứ Nội quy lao động của công ty;\n").italic = True
        p_base.add_run("- Căn cứ Biên bản họp xử lý kỷ luật lao động lập ngày ... ;").italic = True

        # Body - Decisions
        doc.add_paragraph().add_run("GIÁM ĐỐC CÔNG TY QUYẾT ĐỊNH:").bold = True

        # Article 1: Dismissal info
        emp_name = data.get("employee_name", "[Tên người lao động]").upper()
        cccd = data.get("cccd", "[Số CCCD]")
        position = data.get("position", "[Chức vụ]")
        reason = data.get("reason", "[Lý do sa thải]")
        
        p_art1 = doc.add_paragraph()
        p_art1.add_run("Điều 1: ").bold = True
        p_art1.add_run(f"Thi hành kỷ luật lao động bằng hình thức ")
        p_art1.add_run("SA THẢI ").bold = True
        p_art1.add_run(f"đối với ông/bà:\n")
        p_art1.add_run(f"- Họ và tên: ").bold = True
        p_art1.add_run(f"{emp_name}\n")
        p_art1.add_run(f"- Số CCCD: ").bold = True
        p_art1.add_run(f"{cccd}\n")
        p_art1.add_run(f"- Chức vụ/Bộ phận: ").bold = True
        p_art1.add_run(f"{position}\n")
        p_art1.add_run(f"- Lý do sa thải: ").bold = True
        p_art1.add_run(f"{reason}")

        # Article 2: Effective date
        effective_date = data.get("effective_date", "kể từ ngày ký")
        p_art2 = doc.add_paragraph()
        p_art2.add_run("Điều 2: ").bold = True
        p_art2.add_run(f"Quyết định này có hiệu lực kể từ {effective_date}. Mọi chế độ lương, thưởng, bảo hiểm của ông/bà {emp_name} sẽ được giải quyết theo quy định của pháp luật hiện hành và nội quy công ty trong thời gian quy định.")

        # Article 3: Execution responsibility
        p_art3 = doc.add_paragraph()
        p_art3.add_run("Điều 3: ").bold = True
        p_art3.add_run("Ban Giám đốc, bộ phận Nhân sự, bộ phận Kế toán và ông/bà ")
        p_art3.add_run(f"{emp_name} ").bold = True
        p_art3.add_run("chịu trách nhiệm thi hành quyết định này.")

        # Signatures
        p_sign = doc.add_paragraph()
        p_sign.paragraph_format.space_before = Pt(30)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        
        # Left column (Receivers)
        cell_left = table.cell(0, 0)
        p_rec = cell_left.paragraphs[0]
        p_rec.add_run("Nơi nhận:\n").bold = True
        p_rec.add_run(f"- Như Điều 3;\n- Đại diện Công đoàn cơ sở;\n- Lưu: HSNS, VP.")
        p_rec.style.font.size = Pt(10)

        # Right column (Director signature)
        cell_right = table.cell(0, 1)
        p_rep = cell_right.paragraphs[0]
        p_rep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rep.add_run("ĐẠI DIỆN DOANH NGHIỆP\n").bold = True
        p_rep.add_run("GIÁM ĐỐC\n").bold = True
        p_rep.add_run("(Ký tên và đóng dấu)").italic = True

        # Save and return filename
        filename = f"Quyet_dinh_sa_thai_{emp_name.replace(' ', '_')}_{datetime.now().strftime('%H%M%S')}.docx"
        file_path = os.path.join(self.output_dir, filename)
        doc.save(file_path)
        
        return filename

    def generate_nda(self, data: Dict[str, Any]) -> str:
        """Generates a Non-Disclosure Agreement (NDA) document (.docx)."""
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(13)

        # Header
        p_header = doc.add_paragraph()
        p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_header.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
        run.bold = True
        run2 = p_header.add_run("Độc lập - Tự do - Hạnh phúc\n")
        run2.bold = True
        p_header.add_run("---------------o0o---------------")

        # Title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(18)
        p_title.paragraph_format.space_after = Pt(18)
        run_title = p_title.add_run("THỎA THUẬN BẢO MẬT THÔNG TIN (NDA)\n")
        run_title.bold = True
        run_title.font.size = Pt(16)

        # Parties Intro
        p_intro = doc.add_paragraph()
        p_intro.add_run("Thỏa thuận Bảo mật Thông tin này được lập ngày ... giữa hai bên:\n").italic = True
        
        company = data.get("company_name", "[Tên công ty của bạn]").upper()
        rep_name = data.get("rep_name", "[Tên người đại diện]")
        rep_pos = data.get("rep_pos", "[Chức vụ]")
        emp_name = data.get("employee_name", "[Tên nhân viên]")
        cccd = data.get("cccd", "[Số CCCD]")
        
        p_sideA = doc.add_paragraph()
        p_sideA.add_run("Bên A: ").bold = True
        p_sideA.add_run(f"{company}\n")
        p_sideA.add_run(f"- Người đại diện: {rep_name} - Chức vụ: {rep_pos}\n")
        
        p_sideB = doc.add_paragraph()
        p_sideB.add_run("Bên B (Người lao động): ").bold = True
        p_sideB.add_run(f"{emp_name}\n")
        p_sideB.add_run(f"- Số CCCD: {cccd}\n")

        p_agree = doc.add_paragraph()
        p_agree.add_run("Hai bên thống nhất ký kết các điều khoản bảo mật sau đây:")

        # Terms
        p_terms = doc.add_paragraph()
        p_terms.add_run("Điều 1: Thông tin bảo mật\n").bold = True
        p_terms.add_run("Thông tin bảo mật bao gồm toàn bộ dữ liệu khách hàng, mã nguồn phần mềm, chiến lược kinh doanh, bí mật công nghệ, và tài liệu nội bộ của Bên A mà Bên B tiếp cận trong quá trình làm việc.\n\n")
        
        p_terms.add_run("Điều 2: Nghĩa vụ của Bên B\n").bold = True
        p_terms.add_run("1. Không được sao chép, cung cấp hoặc tiết lộ bất kỳ thông tin bảo mật nào cho bên thứ ba mà không có sự đồng ý bằng văn bản của Bên A.\n")
        p_terms.add_run("2. Không sử dụng thông tin bảo mật để trục lợi cá nhân hoặc thành lập doanh nghiệp đối thủ.\n")
        p_terms.add_run("3. Khi chấm dứt hợp đồng lao động, Bên B phải bàn giao lại toàn bộ tài liệu, dữ liệu và thiết bị liên quan đến thông tin bảo mật.\n\n")

        p_terms.add_run("Điều 3: Phạt vi phạm và bồi thường\n").bold = True
        p_terms.add_run("Trong trường hợp Bên B vi phạm Thỏa thuận này, Bên B phải:\n")
        p_terms.add_run("- Chịu mức phạt vi phạm hành chính tương đương 03 tháng lương thực nhận.\n")
        p_terms.add_run("- Bồi thường toàn bộ thiệt hại thực tế phát sinh cho Bên A theo quy định pháp luật sở hữu trí tuệ và thương mại.")

        # Signatures
        p_sign = doc.add_paragraph()
        p_sign.paragraph_format.space_before = Pt(30)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        
        cell_left = table.cell(0, 0)
        p_left = cell_left.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_left.add_run("ĐẠI DIỆN BÊN A\n").bold = True
        p_left.add_run("(Ký tên, đóng dấu)")

        cell_right = table.cell(0, 1)
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_right.add_run("ĐẠI DIỆN BÊN B\n").bold = True
        p_right.add_run("(Ký tên, ghi rõ họ tên)")

        filename = f"Thoa_thuan_bao_mat_NDA_{emp_name.replace(' ', '_')}_{datetime.now().strftime('%H%M%S')}.docx"
        file_path = os.path.join(self.output_dir, filename)
        doc.save(file_path)
        
        return filename

# Quick test if run directly
if __name__ == "__main__":
    gen = LegalDocumentGenerator()
    fn = gen.generate_dismissal_decision({
        "company_name": "Công ty Cổ phần TechSME",
        "employee_name": "Nguyễn Văn A",
        "cccd": "012345678901",
        "position": "Nhân viên Lập trình",
        "reason": "Tự ý nghỉ việc 06 ngày làm việc liên tục không có lý do chính đáng.",
        "effective_date": "ngày 28 tháng 05 năm 2026"
    })
    print("Generated decision:", fn)
